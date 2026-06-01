import argparse
import ast
import datetime as dt
import json
import os
import re
import shutil
import subprocess
import sys
import unicodedata
import urllib.error
import urllib.request
from pathlib import Path


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TOOL_DIR = Path(__file__).resolve().parent
INPUT_DIR = TOOL_DIR / "input"
OUTPUT_DIR = TOOL_DIR / "output"
BACKUP_DIR = TOOL_DIR / "backups"

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".py",
    ".html",
    ".css",
    ".csv",
    ".json",
    ".xml",
    ".sql",
}

ASSET_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".ico",
}

IMPORT_EXTENSIONS = {
    ".xlsx",
    ".xlsm",
    ".csv",
}

GENERATION_ORDER = [
    "core/models.py",
    "core/forms.py",
    "core/views.py",
    "core/templates/core",
    "config/urls.py",
    "core/admin.py",
    "core/permissions.py",
    "core/context_processors.py",
    "core/management/commands/import_data.py",
    "static/css/style.css",
]

CORE_GENERATION_FILES = [
    "core/models.py",
    "core/forms.py",
    "core/views.py",
    "config/urls.py",
    "core/management/commands/import_data.py",
]

MODEL_FIELD_TYPES = (
    "AutoField",
    "BigAutoField",
    "BooleanField",
    "CharField",
    "DateField",
    "DateTimeField",
    "DecimalField",
    "EmailField",
    "FileField",
    "FloatField",
    "ForeignKey",
    "ImageField",
    "IntegerField",
    "OneToOneField",
    "PositiveIntegerField",
    "TextField",
)

PROJECT_FILES = [
    "core/models.py",
    "core/forms.py",
    "core/views.py",
    "core/admin.py",
    "core/permissions.py",
    "core/context_processors.py",
    "core/management/commands/import_data.py",
    "config/urls.py",
    "static/css/style.css",
]

PROJECT_FOLDERS = [
    "core/templates/core",
]

ALLOWED_PATHS = PROJECT_FILES + PROJECT_FOLDERS

SCHEMA_STAGE_FILES = [
    "core/models.py",
    "core/admin.py",
    "core/management/commands/import_data.py",
]

INTERFACE_STAGE_BASE_FILES = [
    "core/forms.py",
    "core/views.py",
    "config/urls.py",
    "core/permissions.py",
    "core/context_processors.py",
    "core/templates/core/base.html",
    "core/templates/core/login.html",
    "static/css/style.css",
]

BLOCKED_GENERATED_SUFFIXES = {".md"}
BLOCKED_GENERATED_PREFIXES = ("docs/",)
BLOCKED_GENERATED_FILES = {"README.md"}

PROJECT_CONTEXT_FILES = [
    "core/models.py",
    "core/forms.py",
    "core/views.py",
    "core/admin.py",
    "core/permissions.py",
    "core/context_processors.py",
    "core/management/commands/import_data.py",
    "config/settings.py",
    "config/urls.py",
    "static/css/style.css",
]

PROJECT_CONTEXT_FOLDERS = [
    "core/templates/core",
]

PROJECT_SHELL_FILES = [
    "config/settings.py",
    "core/templates/core/base.html",
    "core/templates/core/login.html",
    "static/css/style.css",
]

IGNORED_INPUT_FILES = {
    "PUT_EXAM_FILES_HERE.md",
}

ALLOWED_ASSET_TARGETS = [
    "static/images",
    "media/products",
    "core/import",
]

IMPORT_TARGET_DIR = "core/import"
MAX_REPAIR_ATTEMPTS = 5
MAX_FILE_GENERATION_ATTEMPTS = 5

ASSET_TARGET_REPLACEMENTS = {
    "static/icons": "static/images",
    "static/img": "static/images",
    "assets/images": "static/images",
    "assets/icons": "static/images",
}


def load_local_config():
    path = TOOL_DIR / "local_config.json"
    if not path.exists():
        return {}
    return json.loads(read_text(path))


def ensure_dirs():
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)


def log(message, percent=None):
    now = dt.datetime.now().strftime("%H:%M:%S")
    if percent is None:
        print(f"[{now}] {message}")
    else:
        print(f"[{now}] [{percent:3d}%] {message}")


def rel(path):
    return path.resolve().relative_to(PROJECT_ROOT).as_posix()


def normalize_match(value):
    text = str(value or "").replace("\\", "/").strip()
    return unicodedata.normalize("NFC", text).casefold()


def normalize_relative_path(value):
    return str(value or "").replace("\\", "/").strip().lstrip("/")


def is_blocked_generated_path(path_name):
    normalized = normalize_relative_path(path_name)
    suffix = Path(normalized).suffix.lower()
    if suffix in BLOCKED_GENERATED_SUFFIXES:
        return True
    if normalized in BLOCKED_GENERATED_FILES:
        return True
    if normalized == "docs":
        return True
    return any(normalized.startswith(prefix) for prefix in BLOCKED_GENERATED_PREFIXES)


def normalize_asset_target(value):
    target = normalize_relative_path(value)
    for old, new in ASSET_TARGET_REPLACEMENTS.items():
        if target == old:
            return new
        if target.startswith(f"{old}/"):
            return f"{new}/{target[len(old) + 1:]}"
    return target


def is_inside(base, path):
    try:
        path.resolve().relative_to(base.resolve())
        return True
    except ValueError:
        return False


def read_text(path):
    for encoding in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            pass
    return path.read_text(encoding="utf-8", errors="replace")


def read_xlsx(path):
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets[:8]:
        parts.append(f"Лист: {sheet.title}")
        parts.append(f"Размер: строк {sheet.max_row}, столбцов {sheet.max_column}")
        for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_number > 120:
                parts.append("...")
                break
            values = ["" if value is None else str(value) for value in row[:35]]
            parts.append(f"{row_number}: " + "\t".join(values))
    return "\n".join(parts)


def summarize_xlsx(path):
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    parts = [f"Файл: {path.relative_to(INPUT_DIR).as_posix()}"]
    for sheet in workbook.worksheets[:8]:
        rows = sheet.iter_rows(values_only=True)
        first_rows = []
        for row in rows:
            values = ["" if value is None else str(value).strip() for value in row[:35]]
            if any(values):
                first_rows.append(values)
            if len(first_rows) >= 8:
                break

        parts.append(f"Лист: {sheet.title}")
        parts.append(f"Размер: строк {sheet.max_row}, столбцов {sheet.max_column}")
        if first_rows:
            parts.append("Предполагаемые колонки:")
            parts.append(" | ".join(first_rows[0]))
            if len(first_rows) > 1:
                parts.append("Примеры строк:")
                for index, row in enumerate(first_rows[1:], start=2):
                    parts.append(f"{index}: " + " | ".join(row))
        else:
            parts.append("Лист пустой")
    return "\n".join(parts)


def collect_excel_overview():
    files = sorted(
        path
        for path in INPUT_DIR.rglob("*")
        if path.is_file() and path.suffix.lower() in (".xlsx", ".xlsm")
    )
    if not files:
        return "Excel-файлы в tools/input не найдены."

    parts = [
        "# ТАБЛИЦЫ EXCEL ДЛЯ ПРОЕКТИРОВАНИЯ БД",
        "По этим таблицам нужно понять сущности, поля, справочники и связи. Первая строка обычно является заголовками колонок.",
    ]
    for path in files:
        parts.append(f"\n\n===== EXCEL ОБЗОР: {path.relative_to(INPUT_DIR).as_posix()} =====")
        parts.append(summarize_xlsx(path))
    return "\n".join(parts)


def excel_columns(path):
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    sheet = workbook.active
    for row in sheet.iter_rows(values_only=True):
        values = ["" if value is None else str(value).strip() for value in row]
        if any(values):
            return [value for value in values if value]
    return []


def collect_excel_file_info():
    files = sorted(
        path
        for path in INPUT_DIR.rglob("*")
        if path.is_file() and path.suffix.lower() in (".xlsx", ".xlsm")
    )
    result = []
    for path in files:
        result.append(
            {
                "file": path.relative_to(INPUT_DIR).as_posix(),
                "columns": excel_columns(path),
            }
        )
    return result


def collect_import_resource_targets():
    files = sorted(
        path
        for path in INPUT_DIR.rglob("*")
        if path.is_file() and path.suffix.lower() in IMPORT_EXTENSIONS
    )
    if not files:
        return "Файлы Excel/CSV для импорта не найдены."

    parts = [
        "Файлы Excel/CSV будут автоматически скопированы в проект перед проверкой:",
    ]
    for path in files:
        parts.append(f"- {path.relative_to(INPUT_DIR).as_posix()} -> {IMPORT_TARGET_DIR}/{path.name}")
    return "\n".join(parts)


def read_docx(path):
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:10]:
        for row in table.rows[:80]:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def read_pdf(path):
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts = []
    for index, page in enumerate(reader.pages[:25], start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"Страница {index}\n{text}")
    return "\n\n".join(parts)


def read_known_file(path):
    suffix = path.suffix.lower()
    try:
        if suffix in TEXT_EXTENSIONS:
            return read_text(path)
        if suffix in ASSET_EXTENSIONS:
            return f"Изображение. Имя файла: {path.name}. Размер: {path.stat().st_size} байт. Содержимое картинки не анализируется."
        if suffix in (".xlsx", ".xlsm"):
            return read_xlsx(path)
        if suffix == ".docx":
            return read_docx(path)
        if suffix == ".pdf":
            return read_pdf(path)
    except Exception as error:
        return f"Не удалось прочитать файл: {error}"
    return f"Файл не прочитан как текст. Расширение: {suffix}"


def collect_input_files():
    return sorted(
        path
        for path in INPUT_DIR.rglob("*")
        if path.is_file() and path.name not in IGNORED_INPUT_FILES
    )


def collect_assignment():
    files = collect_input_files()
    if not files:
        return "В папке input пока нет файлов задания."
    parts = []

    if any(path.suffix.lower() in (".xlsx", ".xlsm") for path in files):
        parts.append(collect_excel_overview())

    for path in files:
        parts.append(f"\n\n===== ФАЙЛ ЗАДАНИЯ: {path.relative_to(INPUT_DIR).as_posix()} =====")
        parts.append(read_known_file(path)[:30000])
    return "\n".join(parts)


def collect_project():
    paths = []
    for name in PROJECT_CONTEXT_FILES:
        path = PROJECT_ROOT / name
        if path.exists():
            paths.append(path)
    for folder in PROJECT_CONTEXT_FOLDERS:
        base = PROJECT_ROOT / folder
        if base.exists():
            paths.extend(
                path
                for path in base.rglob("*")
                if path.is_file()
                and "__pycache__" not in path.parts
                and path.suffix.lower() in TEXT_EXTENSIONS
            )
    parts = []
    for path in sorted(set(paths)):
        parts.append(f"\n\n===== ФАЙЛ ПРОЕКТА: {rel(path)} =====")
        parts.append(read_known_file(path)[:25000])
    return "\n".join(parts)


def collect_project_shell():
    parts = []
    for name in PROJECT_SHELL_FILES:
        path = PROJECT_ROOT / name
        if path.exists():
            parts.append(f"\n\n===== ФАЙЛ ОБОЛОЧКИ ПРОЕКТА: {rel(path)} =====")
            parts.append(read_known_file(path)[:12000])
    return "\n".join(parts) or "Файлы оболочки проекта не найдены."


def build_context():
    return (
        "# ВЫДАННОЕ ЗАДАНИЕ И ФАЙЛЫ\n"
        f"{collect_assignment()}\n\n"
        "# ТЕКУЩИЕ ФАЙЛЫ ПРОЕКТА\n"
        f"{collect_project()}"
    )


def build_runtime_context(base_context, generated_files=None, extra_notes=""):
    generated_files = generated_files or []
    if not generated_files and not extra_notes:
        return base_context

    parts = [base_context]
    if generated_files:
        parts.append("\n\n# УЖЕ СГЕНЕРИРОВАННЫЕ ФАЙЛЫ ЭТОГО APPLY")
        for item in generated_files:
            parts.append(f"\n\n===== СГЕНЕРИРОВАНО: {item['path']} =====")
            parts.append(item["content"][:25000])
    if extra_notes:
        parts.append("\n\n# ДОПОЛНИТЕЛЬНЫЕ ПРОВЕРКИ И ОШИБКИ")
        parts.append(extra_notes)
    return "\n".join(parts)


def build_blueprint_prompt(assignment_context):
    return f"""
Ты архитектор простого Django-проекта для экзамена.
Нужно внимательно прочитать задание и Excel-таблицы, затем составить единый план проекта.

Главная цель: не смешивать старый шаблон с новым заданием.
Excel-файлы являются главным источником сущностей, полей и связей.
Создавай только те сущности, страницы и импорты, которые нужны по заданию.
Не создавай CRUD для справочников, если задание явно этого не требует.
Главная list-страница должна быть полноценной: поиск по текстовым полям, фильтр по справочнику или статусу, сортировка через GET-параметры.
CRUD-страницы не должны заменять список. Список, добавление, редактирование, просмотр деталей/истории должны быть отдельными страницами, если они есть в задании.
Технические имена классов, функций, url_name и файлов пиши латиницей.
Все видимые пользователю тексты интерфейса, сообщения и заголовки пиши на русском.
Assets клади только в static/images, media/products или core/import. Не используй static/icons.

Ответ верни строго в JSON без Markdown:
{{
  "project_title": "название проекта",
  "main_entity": "главная сущность интерфейса",
  "entities": [
    {{
      "name": "Partner",
      "fields": ["name", "rating"],
      "relations": ["type -> PartnerType"]
    }}
  ],
  "excel_files": [
    {{
      "file": "Partners_import.xlsx",
      "model": "Partner",
      "columns": ["Тип партнера", "Наименование партнера"]
    }}
  ],
  "pages": [
    {{
      "url_name": "partner_list",
      "path": "partners/",
      "view": "PartnerListView",
      "template": "core/templates/core/partner_list.html",
      "purpose": "список партнеров с поиском, фильтром и сортировкой"
    }}
  ],
  "forms": [
    {{
      "name": "PartnerForm",
      "model": "Partner",
      "fields": ["name", "rating"]
    }}
  ],
  "assets": [
    {{
      "source": "logo.png",
      "target": "static/images/logo.png"
    }}
  ],
  "must_have": ["что обязательно проверить"]
}}

{assignment_context}
""".strip()


def build_blueprint_retry_prompt(assignment_context, previous_answer):
    return f"""
Нужно вернуть только JSON с blueprint простого Django-проекта.
Предыдущий ответ не подошел, потому что это был не JSON.

Это учебная задача по Django и Excel-импорту. Никаких секретных данных, вредного кода или обхода защиты не требуется.

Верни строго JSON без Markdown по формату:
{{
  "project_title": "название проекта",
  "main_entity": "главная сущность",
  "entities": [],
  "excel_files": [],
  "pages": [],
  "forms": [],
  "assets": [],
  "must_have": []
}}

Предыдущий ответ:
{previous_answer[:1000]}

Задание:
{assignment_context}
""".strip()


def fallback_blueprint():
    excel_files = collect_excel_file_info()
    main_file = next((item for item in excel_files if item["columns"]), excel_files[0] if excel_files else None)
    main_entity = "Record"
    if main_file:
        stem = Path(main_file["file"]).stem.replace("_import", "").replace(" import", "")
        parts = [part for part in re.split(r"[^A-Za-zА-Яа-я0-9]+", stem) if part]
        if parts:
            main_entity = "".join(part[:1].upper() + part[1:] for part in parts)

    return {
        "project_title": "Учебный Django-проект",
        "main_entity": main_entity,
        "entities": [
            {
                "name": main_entity,
                "fields": main_file["columns"] if main_file else [],
                "relations": [],
            }
        ],
        "excel_files": [{"file": item["file"], "model": "", "columns": item["columns"]} for item in excel_files],
        "pages": [
            {
                "url_name": "record_list",
                "path": "",
                "view": "RecordListView",
                "template": "core/record_list.html",
                "purpose": "список основных записей",
            }
        ],
        "forms": [],
        "assets": [],
        "must_have": ["проект должен проходить python manage.py check"],
    }


def normalize_blueprint(data):
    for asset in data.get("assets", []):
        if "target" in asset:
            asset["target"] = normalize_asset_target(asset["target"])
    return data


def get_blueprint(assignment_context):
    answer = call_yandex(build_blueprint_prompt(assignment_context))
    save_output("blueprint_raw.md", answer)
    try:
        data = extract_json(answer)
    except ValueError:
        answer = call_yandex(build_blueprint_retry_prompt(assignment_context, answer))
        save_output("blueprint_retry_raw.md", answer)
        try:
            data = extract_json(answer)
        except ValueError:
            data = fallback_blueprint()

    data = normalize_blueprint(data)
    text = json.dumps(data, ensure_ascii=False, indent=2)
    save_output("blueprint.md", text)
    return text


def build_generation_context(blueprint):
    return "\n".join(
        [
            "# BLUEPRINT ПРОЕКТА",
            blueprint,
            "\n# ЗАДАНИЕ И EXCEL",
            collect_assignment(),
            "\n# КУДА БУДУТ СКОПИРОВАНЫ ФАЙЛЫ ДЛЯ ИМПОРТА",
            collect_import_resource_targets(),
            "\n# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА",
            collect_project_shell(),
        ]
    )


def find_generated_content(generated_files, path_name):
    for item in reversed(generated_files or []):
        if item["path"] == path_name:
            return item["content"]
    path = PROJECT_ROOT / path_name
    if path.exists() and path.is_file():
        return read_text(path)
    return ""


def has_generated_file(generated_files, path_name):
    return any(item["path"] == path_name for item in generated_files or [])


def generated_template_list(generated_files):
    paths = []
    for item in generated_files or []:
        path_name = item["path"]
        if path_name.startswith("core/templates/core/") and path_name.endswith(".html"):
            paths.append(path_name)
    return "\n".join(f"- {path}" for path in sorted(paths)) or "Шаблоны пока не сгенерированы."


def generated_template_url_names(generated_files):
    result = set()
    for item in generated_files or []:
        path_name = item["path"]
        if path_name.startswith("core/templates/core/") and path_name.endswith(".html"):
            result.update(extract_url_names_from_templates(item["content"]))
    return sorted(result)


def generated_template_url_note(generated_files):
    names = generated_template_url_names(generated_files)
    if not names:
        return "Сгенерированные шаблоны пока не используют url names."
    return "\n".join(f"- {name}" for name in names)


def build_file_context(base_context, generated_files, path_name, extra_notes=""):
    models = find_generated_content(generated_files, "core/models.py")
    forms = find_generated_content(generated_files, "core/forms.py")
    views = find_generated_content(generated_files, "core/views.py")
    urls = find_generated_content(generated_files, "config/urls.py")
    templates = generated_template_list(generated_files)
    template_url_names = generated_template_url_note(generated_files)
    if path_name.startswith("core/templates/core/") and not has_generated_file(generated_files, "config/urls.py"):
        urls = "config/urls.py будет сгенерирован после шаблонов. Используй url_name из BLUEPRINT проекта, а не старые маршруты текущего шаблона."

    if path_name == "core/management/commands/import_data.py":
        return "\n".join(
            [
                "# КОМПАКТНЫЙ КОНТЕКСТ ДЛЯ IMPORT_DATA",
                base_context.split("# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА", 1)[0],
                "Текущий проект является шаблоном. Не смешивай старую предметную область с новой.",
                "Используй только Excel-файлы из блока ниже. Не добавляй импорт старых файлов, которых нет в Excel-обзоре.",
                "Команда импорта должна быть короткой, простой и идемпотентной: get_or_create/update_or_create.",
                "Если модель/поле не описаны в core/models.py ниже, не используй их.",
                "Не импортируй пользователей, заказы, пункты выдачи или старые товары, если таких Excel-файлов нет в новом задании.",
                collect_excel_overview(),
                "\n# ФАЙЛЫ ДЛЯ ИМПОРТА",
                collect_import_resource_targets(),
                "В import_data.py читай Excel/CSV только из папки core/import через settings.BASE_DIR или Path(__file__).resolve().parents. Не используй пути вида 'Ресурсы/...'.",
                "Пример пути: Path(settings.BASE_DIR) / 'core' / 'import' / 'Partners_import.xlsx'.",
                "Если запись уже должна существовать из предыдущего файла, используй get(), а не get_or_create() с неполными обязательными полями.",
                "Порядок импорта: сначала справочники, потом основные сущности, потом таблицы связей/истории.",
                "\n# СГЕНЕРИРОВАННЫЙ core/models.py",
                models,
                "\n# ОШИБКИ/ЗАМЕТКИ",
                extra_notes,
            ]
        )

    if path_name == "core/forms.py":
        return "\n".join(
            [
                "# КОМПАКТНЫЙ КОНТЕКСТ ДЛЯ FORMS",
                base_context.split("# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА", 1)[0],
                "\n# СГЕНЕРИРОВАННЫЙ core/models.py",
                models,
                "\n# ОШИБКИ/ЗАМЕТКИ",
                extra_notes,
            ]
        )

    if path_name == "core/views.py":
        return "\n".join(
            [
                "# КОМПАКТНЫЙ КОНТЕКСТ ДЛЯ VIEWS",
                "Не сохраняй старые view из шаблона, если они не требуются новым заданием.",
                "Создавай только view, перечисленные в BLUEPRINT. Не создавай страницы справочников на всякий случай.",
                "Главная ListView не должна быть пустой. В ней нужен get_queryset() с поиском, фильтрацией и сортировкой через request.GET.",
                "Для поиска используй Q по главным текстовым полям. Для фильтра используй внешний ключ/тип/статус, если такая связь есть. Для сортировки используй белый список разрешенных полей.",
                "Сохраняй все страницы из BLUEPRINT: список, добавление, редактирование, просмотр, историю. Не удаляй функциональность ради упрощения.",
                base_context.split("# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА", 1)[0],
                "\n# СГЕНЕРИРОВАННЫЙ core/models.py",
                models,
                "\n# СГЕНЕРИРОВАННЫЙ core/forms.py",
                forms,
                "\n# ПЛАНИРУЕМЫЕ/СГЕНЕРИРОВАННЫЕ ШАБЛОНЫ",
                templates,
                "\n# ОШИБКИ/ЗАМЕТКИ",
                extra_notes,
            ]
        )

    if path_name == "config/urls.py":
        return "\n".join(
            [
                "# КОМПАКТНЫЙ КОНТЕКСТ ДЛЯ URLS",
                "Импортируй только те view, которые реально есть в core/views.py.",
                "Сохрани маршрут path('admin/', admin.site.urls), если django.contrib.admin есть в INSTALLED_APPS.",
                "Добавь главную страницу '' на основной list-view, чтобы корень сайта открывался.",
                "Добавь login/logout, если base.html или login.html их используют.",
                "Не удаляй url names, которые уже используют сгенерированные шаблоны.",
                "\n# URL NAMES ИЗ СГЕНЕРИРОВАННЫХ ШАБЛОНОВ",
                template_url_names,
                base_context.split("# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА", 1)[0],
                "\n# СГЕНЕРИРОВАННЫЙ core/views.py",
                views,
                "\n# ОШИБКИ/ЗАМЕТКИ",
                extra_notes,
            ]
        )

    if path_name.startswith("core/templates/core/"):
        return "\n".join(
            [
                "# КОМПАКТНЫЙ КОНТЕКСТ ДЛЯ HTML",
                "Не добавляй markdown-ограждения ```html и ``` в файл. Верни чистый Django HTML.",
                "Если в шаблоне есть {% static ... %}, в начале файла обязательно должна быть строка {% load static %}.",
                "Используй только url name из config/urls.py или BLUEPRINT проекта.",
                "Шаблоны Django должны расширять {% extends \"core/base.html\" %}, если это не сам base.html.",
                "Формы изменения данных всегда должны содержать {% csrf_token %}.",
                "List-шаблон должен содержать GET-форму с поиском, фильтром и сортировкой, если это страница списка.",
                "Не убирай кнопки добавления, редактирования, просмотра деталей/истории, если такие маршруты есть в BLUEPRINT или urls.py.",
                base_context.split("# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА", 1)[0],
                "\n# СГЕНЕРИРОВАННЫЙ config/urls.py",
                urls,
                "\n# СГЕНЕРИРОВАННЫЙ core/views.py",
                views,
                "\n# ОШИБКИ/ЗАМЕТКИ",
                extra_notes,
            ]
        )

    if path_name == "core/models.py":
        return "\n".join(
            [
                "# КОМПАКТНЫЙ КОНТЕКСТ ДЛЯ MODELS",
                "Полностью адаптируй модели под BLUEPRINT и Excel. Не сохраняй старые модели шаблона, если они не нужны новому заданию.",
                "Если в config/settings.py стоит AUTH_USER_MODEL = 'core.User', обязательно оставь в core/models.py class User(AbstractUser).",
                "Если нужна авторизация или роли, не удаляй User/Role из шаблона, а адаптируй их минимально.",
                base_context.split("# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА", 1)[0],
                "\n# ОШИБКИ/ЗАМЕТКИ",
                extra_notes,
            ]
        )

    if path_name == "core/admin.py":
        return "\n".join(
            [
                "# КОМПАКТНЫЙ КОНТЕКСТ ДЛЯ ADMIN",
                "Регистрируй только модели из сгенерированного core/models.py.",
                base_context.split("# ОБОЛОЧКА ТЕКУЩЕГО DJANGO-ПРОЕКТА", 1)[0],
                "\n# СГЕНЕРИРОВАННЫЙ core/models.py",
                models,
                "\n# ОШИБКИ/ЗАМЕТКИ",
                extra_notes,
            ]
        )

    return build_runtime_context(base_context, generated_files, extra_notes)


def build_plan_prompt(context):
    return f"""
Ты помогаешь адаптировать простой Django-проект под новое экзаменационное задание.
Сначала внимательно прочитай задание, Excel/PDF/DOCX и текущий код проекта.

Нужно выдать понятный план изменений на русском языке:
1. какие сущности и связи нужны;
2. какие файлы проекта менять;
3. что менять в models.py, forms.py, views.py, urls.py, templates, import_data.py;
4. какие команды потом запустить;
5. какие риски проверить вручную.
6. кратко объясни простыми словами, почему выбраны такие изменения.

Пиши коротко, но конкретно. Код целиком пока не выдавай.

{context}
""".strip()


def build_apply_prompt(context):
    allowed = "\n".join(f"- {path}" for path in ALLOWED_PATHS)
    asset_targets = "\n".join(f"- {path}" for path in ALLOWED_ASSET_TARGETS)
    return f"""
Ты помогаешь адаптировать простой Django-проект под новое экзаменационное задание.
Нужно сначала выбрать список файлов, которые надо изменить.
Не пиши полный код файлов на этом шаге.

Разрешено переписывать только эти файлы и папки:
{allowed}

Не создавай README.md, docs/*.md и любые другие Markdown-файлы. Нужны только файлы кода, шаблонов, стилей и команды импорта.

Картинки из tools/input можно копировать только в эти папки:
{asset_targets}

Excel/CSV из tools/input не указывай в assets. Скрипт сам скопирует все .xlsx, .xlsm и .csv в папку core/import.
import_data.py должен читать данные только из core/import, а не из tools/input и не из папки Ресурсы.

Для assets лучше указывать source как путь из контекста, например "Ресурсы/Мастер пол.ico".
Если известен только файл, можно указать просто "Мастер пол.ico": скрипт найдет его внутри tools/input.
Не используй static/icons. Иконки и логотипы клади в static/images.

Не создавай миграции. Пользователь сам выполнит makemigrations и migrate.
Не добавляй комментарии в код без необходимости.
Не добавляй комментарии в код вообще, кроме стандартных комментариев Django, если они уже есть в файле.
Не меняй секреты, пароли и API-ключи.
Не меняй DATABASES, SECRET_KEY, DEBUG, ALLOWED_HOSTS и настройки подключения к PostgreSQL.
config/settings.py запрещен для записи. Не включай его в files. Пользователь сам настраивает БД и базовые настройки.
Не используй сложные универсальные конфиги.
Сохраняй стиль простого Django-кода: models, forms, views, templates, management command.
Если новое задание меняет предметную область, замени старые сущности шаблона на новые. Не смешивай старую и новую предметные области.
Excel-файлы из tools/input являются главным источником для моделей и импорта.
Технические имена в коде пиши латиницей, видимые пользователю тексты интерфейса пиши на русском.
Главная list-страница обязана иметь поиск, фильтр и сортировку через GET-параметры. Не заменяй ее пустым ListView.
Если в задании есть CRUD, не удаляй add/update/detail/history страницы ради упрощения.
Не создавай страницы для Supplier, Manufacturer, Category, Unit, PickupPoint, OrderStatus, Product, Order, если они не указаны в BLUEPRINT проекта.
Обязательно включи core/templates/core/base.html, если меняешь urls.py, views.py или основные шаблоны.
Обязательно включи core/templates/core/login.html, если в base.html или urls.py есть вход пользователя.
config/urls.py должен сохранять admin/ и иметь главную страницу '' на основной список.

Ответ верни строго в JSON без Markdown:
{{
  "summary": "кратко что изменено",
  "steps": ["что потом проверить"],
  "commands": ["команды после применения"],
  "files": [
    {{
      "path": "core/models.py",
      "reason": "зачем менять этот файл"
    }}
  ],
  "assets": [
    {{
      "source": "images/logo.png",
      "target": "static/images/logo.png"
    }}
  ]
}}

В files включай только те файлы, которые реально надо заменить.
В files нельзя указывать папку. Если надо изменить шаблоны, указывай конкретные HTML-файлы, например:
- core/templates/core/partner_list.html
- core/templates/core/partner_form.html
- core/templates/core/partner_history.html
В assets включай только картинки, которые нужно скопировать из tools/input в проект. source - путь относительно tools/input. target - путь относительно корня Django-проекта.
Если картинки не нужны, верни пустой список assets.

{context}
""".strip()


def build_folder_files_prompt(context, folder_name, reason):
    return f"""
Ты выбрал папку вместо конкретных файлов. Нужно заменить это на список конкретных файлов.

Папка:
{folder_name}

Зачем нужны файлы:
{reason}

Правила:
- верни только JSON без Markdown;
- в files указывай только конкретные файлы, не папки;
- все пути должны начинаться с "{folder_name.rstrip('/')}/";
- для Django-шаблонов используй расширение .html;
- файлов должно быть немного, только самые необходимые для задания.

Формат:
{{
  "files": [
    {{
      "path": "{folder_name.rstrip('/')}/example.html",
      "reason": "зачем нужен файл"
    }}
  ]
}}

{context}
""".strip()


def build_repair_prompt(context, error_text):
    allowed = "\n".join(f"- {path}" for path in ALLOWED_PATHS)
    return f"""
После автоматического изменения проекта команда проверки Django упала.
Нужно выбрать минимальный список файлов, которые надо исправить.
Не пиши полный код файлов на этом шаге.
Текущий контекст ниже содержит уже записанные файлы после apply и предыдущих исправлений. Исправляй именно текущие файлы, не возвращай старую предметную область шаблона.

Разрешено переписывать только эти файлы и папки:
{allowed}

Правила:
- исправляй причину ошибки, а не только первую строку traceback;
- проверь согласованность urls.py, views.py, forms.py, models.py и templates;
- если в urls.py импортируется класс, он обязан существовать во views.py;
- если view использует template_name, такой шаблон обязан быть создан;
- если форма использует модель и поля, они обязаны существовать в models.py;
- если import_data.py записывает поля модели, эти поля обязаны существовать;
- если ошибка связана с шаблоном Django, убери markdown-ограждения ```html/``` и проверь {{% load static %}};
- если шаблон использует {{% static %}}, в нем должен быть {{% load static %}};
- если шаблон содержит POST-форму, в нем должен быть {{% csrf_token %}};
- если исправляешь import_data.py, Excel/CSV читай только из core/import через BASE_DIR, без путей "Ресурсы/..." и "tools/input/...";
- если исправляешь get_or_create для связанной записи, не создавай объект с неполными обязательными полями, используй get() после импорта справочников;
- config/settings.py запрещен для записи. Не включай его в files. Пользователь сам настраивает БД и базовые настройки;
- если core/views.py содержит list-страницу без поиска/фильтрации/сортировки, исправь views.py и list-шаблон;
- не создавай миграции;
- не меняй секреты и API-ключи;
- не меняй DATABASES, SECRET_KEY, DEBUG, ALLOWED_HOSTS и настройки подключения к PostgreSQL;
- в files указывай только конкретные файлы, не папки.
- не добавляй README.md, docs/*.md и любые другие Markdown-файлы.

Ответ верни строго в JSON без Markdown:
{{
  "summary": "кратко что исправлено",
  "steps": ["что проверить после исправления"],
  "commands": ["команды после исправления"],
  "files": [
    {{
      "path": "core/views.py",
      "reason": "зачем менять этот файл"
    }}
  ],
  "assets": []
}}

Ошибка проверки:
{error_text}

Контекст проекта:
{context}
""".strip()


def build_file_prompt(context, path_name, reason):
    return f"""
Ты переписываешь один файл Django-проекта под экзаменационное задание.
Нужно вернуть полный новый текст только для одного файла.

Файл:
{path_name}

Зачем он меняется:
{reason}

Правила:
- пиши простой понятный Django-код;
- не добавляй лишних абстракций;
- не создавай миграции;
- не меняй API-ключи и пароли;
- не меняй настройки подключения к базе данных;
- не меняй config/settings.py;
- не добавляй комментарии без необходимости;
- не добавляй комментарии в код вообще;
- используй уже сгенерированные файлы из контекста как источник правды;
- технические имена классов, функций, переменных, url_name и файлов пиши латиницей;
- видимые пользователю заголовки, кнопки и сообщения пиши на русском;
- если пишешь config/urls.py, импортируй только те view-классы или функции, которые реально есть в core/views.py;
- если пишешь config/urls.py, сохрани admin/ и добавь главную страницу '' на основной list-view;
- если пишешь core/views.py, template_name должен ссылаться на реально существующий или генерируемый шаблон;
- если пишешь core/views.py и это ListView, обязательно реализуй get_queryset() с поиском, фильтрацией и сортировкой через GET;
- если пишешь core/forms.py, поля формы должны существовать в модели из core/models.py;
- если пишешь import_data.py, не записывай поля, которых нет в моделях;
- если пишешь import_data.py, используй только Excel/CSV-файлы из текущего задания, которые будут лежать в core/import;
- если пишешь import_data.py, не используй пути вида "Ресурсы/file.xlsx", "tools/input/file.xlsx" и абсолютные пути;
- если пишешь import_data.py, для связей используй уже импортированные объекты через get(), а не get_or_create() с неполными обязательными полями;
- если пишешь import_data.py, делай файл идемпотентным через update_or_create/get_or_create и правильный порядок импорта;
- не смешивай старые сущности шаблона с новой предметной областью;
- если это HTML, верни полный HTML-шаблон;
- если это HTML, не добавляй строки ```html и ```;
- если это HTML list-страницы, добавь GET-форму поиска, фильтра и сортировки;
- если это HTML и используется {{% static %}}, добавь {{% load static %}} в начало файла;
- если это HTML с form method="post", обязательно добавь {{% csrf_token %}};
- если это Python, верни полный Python-файл.

Ответ верни строго в таком формате без Markdown:
===FILE:{path_name}===
полный новый текст файла
===END===

Не добавляй текст до `===FILE:{path_name}===` и после `===END===`.

{context}
""".strip()


def build_file_retry_prompt(context, path_name, reason, previous_answer):
    truncated_note = ""
    if "===END===" not in previous_answer:
        truncated_note = (
            "Предыдущий ответ был обрезан или не содержал ===END===. "
            "Сделай файл короче. Убери поддержку старых Excel-файлов и старых сущностей шаблона, если их нет в текущем задании."
        )
    return f"""
Нужно вернуть полный текст файла Django-проекта.
Предыдущий ответ не подошел, потому что в нем не было обязательных маркеров.
{truncated_note}

Файл:
{path_name}

Зачем он меняется:
{reason}

Верни ответ строго в таком формате:
===FILE:{path_name}===
полный текст файла без Markdown
===END===

Не добавляй пояснения. Не используй Markdown. Не отказывайся от ответа: это обычный учебный Django-код.
Если файл большой, выбери минимальное простое решение, которое проходит Django check.

Предыдущий неправильный ответ:
{previous_answer[:2000]}

Контекст:
{context}
""".strip()


def build_consistency_repair_prompt(context, errors):
    allowed = "\n".join(f"- {path}" for path in ALLOWED_PATHS)
    return f"""
Сгенерированные файлы Django-проекта не согласованы между собой.
Нужно выбрать минимальный список файлов, которые надо перегенерировать.
Не пиши полный код файлов на этом шаге.

Разрешено менять только:
{allowed}

Правила:
- если urls.py импортирует отсутствующие view, обычно исправляй urls.py под реальные views, а не добавляй старые классы;
- если view ссылается на отсутствующий шаблон, добавь конкретный HTML-файл или исправь template_name;
- если шаблон использует неизвестный url name, исправь шаблон или urls.py;
- если форма использует поля, которых нет в модели, исправь форму или модель;
- если import_data.py пишет в поля, которых нет в модели, исправь import_data.py или модель;
- если HTML содержит ```html/```, убери эти строки;
- если HTML использует {{% static %}}, добавь {{% load static %}};
- если HTML содержит POST-форму, добавь {{% csrf_token %}};
- если import_data.py читает Excel/CSV из Ресурсы или tools/input, перепиши путь на core/import через BASE_DIR;
- если get_or_create создает связанную запись без обязательных полей, используй get() после импорта справочника;
- если list-страница стала пустым ListView без поиска/фильтрации/сортировки, исправь views.py и HTML-шаблон списка;
- в files указывай только конкретные файлы, не папки;
- не меняй DATABASES, SECRET_KEY, DEBUG, ALLOWED_HOSTS и настройки PostgreSQL.

Ответ верни строго в JSON без Markdown:
{{
  "summary": "кратко что надо исправить",
  "steps": ["что проверить"],
  "commands": ["команды после применения"],
  "files": [
    {{
      "path": "config/urls.py",
      "reason": "зачем перегенерировать этот файл"
    }}
  ],
  "assets": []
}}

Ошибки согласованности:
{errors}

Контекст:
{context}
""".strip()


def collect_output_notes():
    parts = []
    for name in (
        "blueprint.md",
        "plan.md",
        "apply_report.md",
        "apply_manifest_response.md",
        "consistency_errors_final.md",
        "check_after_apply.md",
        "manual_check.md",
    ):
        path = OUTPUT_DIR / name
        if path.exists():
            parts.append(f"\n\n===== ФАЙЛ OUTPUT: {name} =====")
            parts.append(read_text(path)[:20000])
    return "\n".join(parts) or "Файлов с предыдущими ответами пока нет."


def question_file_candidates(question):
    names = set()
    suffixes = "py|html|css|csv|json|xml|sql|txt|md"
    patterns = [
        rf"`([^`]+?\.(?:{suffixes}))`",
        rf'"([^"]+?\.(?:{suffixes}))"',
        rf"'([^']+?\.(?:{suffixes}))'",
        rf"[A-Za-z]:[\\/][^\r\n\"'<>|]*?\.(?:{suffixes})",
        rf"[\wА-Яа-яЁё./\\:-]+?\.(?:{suffixes})",
    ]
    for pattern in patterns:
        for match in re.findall(pattern, question or "", flags=re.IGNORECASE):
            if isinstance(match, tuple):
                match = next((part for part in match if part), "")
            name = match.strip(".,;:!?\"'`()[]{} ")
            if name:
                names.add(name)

    result = []
    for name in sorted(names, key=len, reverse=True):
        normalized = normalize_match(name)
        if any(normalize_match(current).endswith(normalized) for current in result):
            continue
        result.append(name)
    return sorted(result)


def resolve_question_file(name):
    normalized = normalize_relative_path(name)
    candidates = []
    direct = Path(name)
    if direct.is_absolute():
        candidates.append(direct)

    if normalized.startswith("tools/output/"):
        candidates.append(PROJECT_ROOT / normalized)
    elif normalized.startswith("output/"):
        candidates.append(TOOL_DIR / normalized)
    elif "/" in normalized:
        candidates.append(PROJECT_ROOT / normalized)
        candidates.append(TOOL_DIR / normalized)
    else:
        candidates.append(OUTPUT_DIR / normalized)
        candidates.append(PROJECT_ROOT / normalized)
        candidates.extend(OUTPUT_DIR.rglob(normalized))

    for path in candidates:
        try:
            resolved = path.resolve()
        except OSError:
            continue
        if not resolved.exists() or not resolved.is_file():
            continue
        if not (is_inside(PROJECT_ROOT, resolved) or is_inside(TOOL_DIR, resolved)):
            continue
        if resolved.suffix.lower() not in TEXT_EXTENSIONS:
            continue
        return resolved
    return None


def collect_question_files(question):
    parts = []
    seen = set()
    for name in question_file_candidates(question)[:12]:
        path = resolve_question_file(name)
        if not path or path in seen:
            continue
        seen.add(path)
        label = path.relative_to(PROJECT_ROOT).as_posix() if is_inside(PROJECT_ROOT, path) else path.name
        parts.append(f"\n\n===== ФАЙЛ ИЗ ВОПРОСА: {label} =====")
        parts.append(read_text(path)[:30000])
    return "\n".join(parts) or "В вопросе не указаны конкретные файлы или они не найдены."


def build_question_prompt(context, question, history):
    history_text = "\n".join(history[-8:]) if history else "Истории диалога пока нет."
    notes = collect_output_notes()
    question_files = collect_question_files(question)
    return f"""
Ты отвечаешь на вопросы по простому Django-проекту для экзамена.
Отвечай по-русски, коротко и практически.

Правила ответа:
- сначала дай прямой ответ;
- если нужно действие, дай шаги;
- если нужен код, покажи только нужный небольшой фрагмент;
- называй конкретные файлы проекта;
- если вопрос про ошибку, объясни причину и что проверить первым;
- не переписывай весь проект без просьбы.

Контекст задания и проекта:
{context}

Предыдущие ответы инструмента:
{notes}

Файлы, явно названные в вопросе:
{question_files}

История текущего чата:
{history_text}

Вопрос пользователя:
{question}
""".strip()


def build_audit_prompt(context, condition):
    condition_text = condition or "Проверь весь проект на соответствие заданию и типовым требованиям экзамена."
    return f"""
Ты проверяешь Django-проект для экзамена по выданному заданию.
Нужно не писать код, а провести строгий аудит соответствия.

Проверяемое условие:
{condition_text}

Проверь:
1. модели и связи по Excel/заданию;
2. 3НФ и отсутствие лишнего дублирования;
3. импорт данных, идемпотентность, правильные внешние ключи;
4. список: поиск, фильтр, сортировка, пагинация если требуется;
5. CRUD: добавление, редактирование, просмотр, удаление если требуется;
6. url names и все страницы из задания;
7. формы, CSRF, ограничения доступа по ролям;
8. изображения, media/static, Pillow если требуется;
9. команды запуска, миграции и импорт.

Ответ дай по-русски:
- сначала вердикт: готово/частично/не готово;
- затем таблицу "требование - статус - где смотреть - что исправить";
- в конце короткий список действий по приоритету.

Контекст задания и проекта:
{context}
""".strip()


def call_yandex(prompt):
    from yandex_ai_studio_sdk import AIStudio

    local_config = load_local_config()
    api_key = local_config.get("YC_API_KEY")
    folder_id = local_config.get("YC_FOLDER_ID")
    model_name = local_config.get("YC_MODEL", "yandexgpt")
    model_version = local_config.get("YC_MODEL_VERSION", "rc")
    api_mode = local_config.get("YC_API_MODE", "sdk")
    temperature = float(local_config.get("YC_TEMPERATURE", "0.2"))

    if not api_key or not folder_id:
        raise RuntimeError("Заполни YC_API_KEY и YC_FOLDER_ID в tools/local_config.json")
    if "/folders/" in folder_id:
        folder_id = folder_id.rstrip("/").split("/folders/", 1)[1].split("?", 1)[0]

    if api_mode == "openai" or model_name.lower().startswith("qwen") or model_name.startswith("gpt://"):
        if model_version == "rc":
            model_version = "latest"
        return call_yandex_openai(prompt, api_key, folder_id, model_name, model_version, temperature)

    sdk = AIStudio(folder_id=folder_id, auth=api_key)
    try:
        model = sdk.models.completions(model_name, model_version=model_version)
    except TypeError:
        model = sdk.models.completions(model_name)
    model = model.configure(temperature=temperature)
    result = model.run(prompt)
    return result[0].text


def call_yandex_openai(prompt, api_key, folder_id, model_name, model_version, temperature):
    if model_name.startswith("gpt://"):
        model_uri = model_name
    else:
        model_uri = f"gpt://{folder_id}/{model_name}/{model_version}"

    body = json.dumps(
        {
            "model": model_uri,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": temperature,
        }
    ).encode("utf-8")
    request = urllib.request.Request(
        "https://llm.api.cloud.yandex.net/v1/chat/completions",
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "OpenAI-Project": folder_id,
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        details = error.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Ошибка Yandex OpenAI-compatible API: {error.code} {details}") from error

    return data["choices"][0]["message"]["content"]


def extract_json(text):
    cleaned = text.strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)\s*```", cleaned, re.DOTALL)
    if fenced:
        cleaned = fenced.group(1).strip()
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Модель не вернула JSON с объектом.")
    try:
        return json.loads(cleaned[start : end + 1])
    except json.JSONDecodeError as error:
        raise ValueError(
            "вернул невалидный JSON. "
            f"Строка {error.lineno}, столбец {error.colno}. "
            "Сырой ответ сохранен в tools/output/last_response.md"
        ) from error


def extract_file_content(text, path_name):
    start_marker = f"===FILE:{path_name}==="
    end_marker = "===END==="
    start = text.find(start_marker)
    if start == -1:
        fenced = re.findall(r"```(?:python|html|css|text)?\s*(.*?)\s*```", text, re.DOTALL)
        if len(fenced) == 1:
            return fenced[0].strip()
        raise ValueError(f"не вернул начало файла {path_name}. Сырой ответ сохранен в tools/output/last_response.md")
    start += len(start_marker)
    end = text.find(end_marker, start)
    if end == -1:
        raise ValueError(f"не вернул конец файла {path_name}. Сырой ответ сохранен в tools/output/last_response.md")
    return text[start:end].strip()


def remove_python_comments(content):
    lines = []
    for line in content.splitlines():
        stripped = line.lstrip()
        if stripped.startswith("#"):
            continue
        lines.append(line.rstrip())
    return "\n".join(lines).strip() + "\n"


def remove_html_comments(content):
    return re.sub(r"\s*<!--.*?-->\s*", "\n", content, flags=re.DOTALL).strip() + "\n"


def remove_css_comments(content):
    return re.sub(r"/\*.*?\*/", "", content, flags=re.DOTALL).strip() + "\n"


def strip_markdown_fence(content):
    text = content.strip()
    match = re.fullmatch(r"```(?:python|html|css|text)?\s*(.*?)\s*```", text, re.DOTALL)
    if match:
        return match.group(1).strip() + "\n"

    lines = text.splitlines()
    if lines and re.fullmatch(r"```(?:python|html|css|text)?\s*", lines[0].strip()):
        lines = lines[1:]
    if lines and lines[-1].strip() == "```":
        lines = lines[:-1]
    return "\n".join(lines).strip() + "\n"


def ensure_template_static_load(content):
    if "{% static" not in content or "{% load static %}" in content:
        return content
    return "{% load static %}\n" + content.lstrip()


def clean_generated_content(path_name, content):
    content = strip_markdown_fence(content)
    suffix = Path(path_name).suffix.lower()
    if suffix == ".py":
        return remove_python_comments(content)
    if suffix == ".html":
        return ensure_template_static_load(remove_html_comments(content))
    if suffix == ".css":
        return remove_css_comments(content)
    return content.strip() + "\n"


def is_allowed(relative_path):
    normalized = Path(relative_path).as_posix().lstrip("/")
    if is_blocked_generated_path(normalized):
        return False
    if ".." in Path(normalized).parts:
        return False
    for allowed in ALLOWED_PATHS:
        allowed = Path(allowed).as_posix().rstrip("/")
        if normalized == allowed or normalized.startswith(f"{allowed}/"):
            return True
    return False


def is_allowed_folder(relative_path):
    normalized = normalize_relative_path(relative_path).rstrip("/")
    if ".." in Path(normalized).parts:
        return False
    return any(normalized == Path(folder).as_posix().rstrip("/") for folder in PROJECT_FOLDERS)


def looks_like_folder(relative_path):
    normalized = normalize_relative_path(relative_path).rstrip("/")
    if is_allowed_folder(normalized):
        return True
    target = PROJECT_ROOT / normalized
    return target.exists() and target.is_dir()


def validate_file_path(path_name):
    normalized = normalize_relative_path(path_name)
    path = Path(normalized)
    if not normalized or path.is_absolute() or ".." in path.parts:
        raise ValueError(f"Некорректный путь файла: {path_name}")
    if is_blocked_generated_path(normalized):
        raise ValueError(f"Markdown и документация не генерируются: {normalized}")
    if looks_like_folder(normalized):
        raise ValueError(f"Модель указала папку вместо файла: {normalized}")
    if path.suffix.lower() not in TEXT_EXTENSIONS:
        raise ValueError(f"У файла должно быть текстовое расширение: {normalized}")
    if not is_allowed(normalized):
        raise ValueError(f"Запрещенный путь для записи: {normalized}")
    return normalized


def is_asset_target_allowed(relative_path):
    normalized = Path(relative_path).as_posix().lstrip("/")
    if ".." in Path(normalized).parts:
        return False
    for allowed in ALLOWED_ASSET_TARGETS:
        allowed = Path(allowed).as_posix().rstrip("/")
        if normalized.startswith(f"{allowed}/"):
            return True
    return False


def list_input_assets():
    return sorted(
        path
        for path in INPUT_DIR.rglob("*")
        if path.is_file() and path.suffix.lower() in ASSET_EXTENSIONS
    )


def format_asset_options(paths):
    return ", ".join(path.relative_to(INPUT_DIR).as_posix() for path in paths[:12])


def find_input_asset(source_name):
    source_name = normalize_relative_path(source_name)
    source_path = Path(source_name)
    if not source_name or source_path.is_absolute() or ".." in source_path.parts:
        raise ValueError(f"Некорректный путь картинки: {source_name}")

    direct = (INPUT_DIR / source_path).resolve()
    if is_inside(INPUT_DIR, direct) and direct.exists() and direct.is_file():
        return direct

    assets = list_input_assets()
    requested_path = normalize_match(source_name)
    requested_name = normalize_match(source_path.name)
    requested_stem = normalize_match(source_path.stem)

    matches = [
        path
        for path in assets
        if normalize_match(path.relative_to(INPUT_DIR).as_posix()) == requested_path
    ]
    if not matches:
        matches = [path for path in assets if normalize_match(path.name) == requested_name]
    if not matches and source_path.suffix == "":
        matches = [path for path in assets if normalize_match(path.stem) == requested_stem]

    if not matches:
        options = format_asset_options(assets)
        if options:
            raise ValueError(f"Картинка не найдена в input: {source_name}. Доступные картинки: {options}")
        raise ValueError(f"Картинка не найдена в input: {source_name}. В tools/input нет картинок.")

    if len(matches) > 1:
        suffix = source_path.suffix.lower()
        if suffix:
            same_suffix = [path for path in matches if path.suffix.lower() == suffix]
            if len(same_suffix) == 1:
                return same_suffix[0]
            if same_suffix:
                matches = same_suffix
        raise ValueError(
            "Найдено несколько картинок с похожим именем. "
            f"Укажи точный путь из input: {format_asset_options(matches)}"
        )

    return matches[0]


def write_files(data):
    timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_root = BACKUP_DIR / timestamp
    changed = []

    prepared = []
    for item in data.get("files", []):
        path_name = validate_file_path(item.get("path", ""))
        content = item.get("content")
        if not path_name or content is None:
            continue
        target = (PROJECT_ROOT / path_name).resolve()
        if not is_inside(PROJECT_ROOT, target):
            raise ValueError(f"Путь вышел за пределы проекта: {path_name}")
        if target.exists() and target.is_dir():
            raise ValueError(f"Нельзя записать файл поверх папки: {path_name}")
        prepared.append((path_name, content, target))

    for path_name, content, target in prepared:
        if target.exists():
            backup = backup_root / path_name
            backup.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(target, backup)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding="utf-8")
        changed.append(path_name)
    return changed, backup_root


def validate_assets(data):
    assets = []
    for item in data.get("assets", []):
        source_name = normalize_relative_path(item.get("source", ""))
        target_name = normalize_asset_target(item.get("target", ""))
        if not source_name or not target_name:
            continue

        source = find_input_asset(source_name)
        if Path(target_name).suffix == "":
            target_name = f"{target_name.rstrip('/')}/{source.name}"

        target = (PROJECT_ROOT / target_name).resolve()
        if not is_inside(PROJECT_ROOT, target):
            raise ValueError(f"Путь картинки вышел за пределы проекта: {target_name}")
        if source.suffix.lower() not in ASSET_EXTENSIONS:
            raise ValueError(f"Файл не является разрешенной картинкой: {source_name}")
        if not is_asset_target_allowed(target_name):
            raise ValueError(f"Запрещенный путь для картинки: {target_name}")
        assets.append((source, target, target_name))
    return assets


def copy_assets(assets, backup_root):
    copied = []
    for source, target, target_name in assets:
        if target.exists():
            backup = backup_root / target_name
            backup.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(target, backup)
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)
        copied.append(target_name)
    return copied


def list_input_import_resources():
    return sorted(
        path
        for path in INPUT_DIR.rglob("*")
        if path.is_file() and path.suffix.lower() in IMPORT_EXTENSIONS
    )


def copy_import_resources(backup_root):
    copied = []
    target_dir = (PROJECT_ROOT / IMPORT_TARGET_DIR).resolve()
    target_dir.mkdir(parents=True, exist_ok=True)
    for source in list_input_import_resources():
        target_name = f"{IMPORT_TARGET_DIR}/{source.name}"
        target = (PROJECT_ROOT / target_name).resolve()
        if not is_inside(PROJECT_ROOT, target):
            raise ValueError(f"Путь файла импорта вышел за пределы проекта: {target_name}")
        if target.exists():
            backup = backup_root / target_name
            backup.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(target, backup)
        shutil.copy2(source, target)
        copied.append(target_name)
    return copied


def run_project_command(command, timeout=180):
    env = os.environ.copy()
    env.pop("VIRTUAL_ENV", None)
    completed = subprocess.run(
        command,
        cwd=PROJECT_ROOT,
        env=env,
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        timeout=timeout,
    )
    output = "\n".join(part for part in (completed.stdout, completed.stderr) if part.strip())
    return completed.returncode == 0, output.strip()


TEMPLATE_COMPILE_SCRIPT = """
from django.template.loader import get_template
from core import views

templates = {'core/base.html', 'core/login.html'}
for value in vars(views).values():
    if getattr(value, '__module__', '') != 'core.views':
        continue
    template_name = getattr(value, 'template_name', None)
    if isinstance(template_name, str):
        templates.add(template_name)

for template_name in sorted(templates):
    get_template(template_name)
print('Templates OK')
""".strip()


PAGE_SMOKE_SCRIPT = """
from django.apps import apps
from django.db import connection
from django.test import Client
from django.urls import URLPattern, URLResolver, get_resolver

try:
    existing_tables = set(connection.introspection.table_names())
except Exception as error:
    print(f'Page smoke skipped: database is not ready: {error}')
    raise SystemExit(0)

required_tables = {
    model._meta.db_table
    for model in apps.get_app_config('core').get_models()
    if model._meta.managed
}
missing_tables = sorted(required_tables - existing_tables)
if missing_tables:
    print('Page smoke skipped: missing tables: ' + ', '.join(missing_tables[:10]))
    raise SystemExit(0)

def walk(patterns, prefix=''):
    for pattern in patterns:
        if isinstance(pattern, URLPattern):
            route = prefix + str(pattern.pattern)
            if '<' not in route and not route.startswith('admin/'):
                yield '/' + route.lstrip('/')
        elif isinstance(pattern, URLResolver):
            yield from walk(pattern.url_patterns, prefix + str(pattern.pattern))

client = Client()
errors = []
for path in sorted(set(walk(get_resolver().url_patterns))):
    try:
        response = client.get(path)
    except Exception as error:
        errors.append(f'{path}: {error.__class__.__name__}: {error}')
        continue
    print(f'{path}: HTTP {response.status_code}')
    if response.status_code >= 500:
        errors.append(f'{path}: HTTP {response.status_code}')

if errors:
    print('\\n'.join(errors))
    raise SystemExit(1)

print('Page smoke OK')
""".strip()


def validate_project():
    consistency_errors = check_generated_consistency({"files": []})
    if consistency_errors:
        text = "\n".join(f"- {error}" for error in consistency_errors)
        return False, f"$ internal consistency check\n{text}"

    commands = [
        ["uv", "run", "python", "-m", "py_compile", "core/models.py", "core/forms.py", "core/views.py", "core/management/commands/import_data.py", "config/urls.py"],
        ["uv", "run", "python", "manage.py", "check"],
        ["uv", "run", "python", "manage.py", "shell", "-c", TEMPLATE_COMPILE_SCRIPT],
    ]
    report = []
    for command in commands:
        command_text = " ".join(command)
        log(f"Проверяю: {command_text}")
        ok, output = run_project_command(command)
        report.append(f"$ {command_text}\n{output or 'OK'}")
        if not ok:
            return False, "\n\n".join(report)
    return True, "\n\n".join(report)


def validate_pages():
    ok, output = validate_project()
    report = [output]
    if not ok:
        return False, "\n\n".join(report)

    command = ["uv", "run", "python", "manage.py", "shell", "-c", PAGE_SMOKE_SCRIPT]
    command_text = " ".join(command)
    log(f"Проверяю страницы: {command_text}")
    ok, output = run_project_command(command)
    report.append(f"$ {command_text}\n{output or 'OK'}")
    if "Page smoke skipped" in output:
        report.append("База данных еще не готова для проверки страниц. Сначала выполни makemigrations, migrate и import_data.")
        return False, "\n\n".join(report)
    return ok, "\n\n".join(report)


def page_test_waits_for_database(output):
    return "Page smoke skipped" in output or "База данных еще не готова" in output


def safe_output_name(path_name):
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", path_name.replace("/", "_").replace("\\", "_"))


def file_order_key(item):
    path_name = item["path"] if isinstance(item, dict) else item
    for index, prefix in enumerate(GENERATION_ORDER):
        if path_name == prefix or path_name.startswith(f"{prefix}/"):
            return index
    return len(GENERATION_ORDER)


def sort_file_items(files):
    return sorted(files, key=lambda item: (file_order_key(item), item["path"]))


def filter_generated_files(files):
    result = []
    for item in files or []:
        path_name = normalize_relative_path(item.get("path", ""))
        if not path_name:
            continue
        if is_blocked_generated_path(path_name):
            log(f"Пропускаю документацию из генерации: {path_name}")
            continue
        result.append(item)
    return result


def ensure_core_files(manifest):
    files = filter_generated_files(manifest.get("files", []))
    paths = {normalize_relative_path(item.get("path", "")) for item in files}
    required = [
        "core/models.py",
        "core/forms.py",
        "core/views.py",
        "config/urls.py",
        "core/admin.py",
        "core/management/commands/import_data.py",
    ]
    for path_name in required:
        if path_name not in paths:
            files.append(
                {
                    "path": path_name,
                    "reason": "обязательный файл логики проекта, должен быть согласован с новым заданием",
                }
            )
            paths.add(path_name)
    manifest["files"] = files
    return manifest


def ensure_blueprint_templates(manifest, blueprint_text):
    try:
        blueprint = json.loads(blueprint_text)
    except json.JSONDecodeError:
        return manifest

    files = filter_generated_files(manifest.get("files", []))
    paths = {normalize_relative_path(item.get("path", "")) for item in files}
    for page in blueprint.get("pages", []):
        template = normalize_template_file_path(page.get("template", ""))
        if not template:
            continue
        if not template.startswith("core/templates/core/") or not template.endswith(".html"):
            continue
        if template in paths:
            continue
        files.append(
            {
                "path": template,
                "reason": "шаблон страницы из blueprint проекта",
            }
        )
        paths.add(template)

    manifest["files"] = files
    return manifest


def ensure_support_files(manifest):
    files = filter_generated_files(manifest.get("files", []))
    paths = {normalize_relative_path(item.get("path", "")) for item in files}
    touches_interface = any(
        path in paths or any(current.startswith(f"{path}/") for current in paths)
        for path in ("core/views.py", "config/urls.py", "core/templates/core")
    )
    if touches_interface and "core/templates/core/base.html" not in paths:
        files.append(
            {
                "path": "core/templates/core/base.html",
                "reason": "обновить базовый шаблон под новые url names и новую предметную область",
            }
        )
    if touches_interface and "core/templates/core/login.html" not in paths:
        files.append(
            {
                "path": "core/templates/core/login.html",
                "reason": "согласовать страницу входа с новыми маршрутами проекта",
            }
        )
    manifest["files"] = files
    return manifest


def generated_map(data):
    return {item["path"]: item["content"] for item in data.get("files", [])}


def read_project_or_generated(data, path_name):
    current = generated_map(data)
    if path_name in current:
        return current[path_name]
    path = PROJECT_ROOT / path_name
    if path.exists() and path.is_file():
        return read_text(path)
    return ""


def generated_template_paths(data):
    paths = set()
    for item in data.get("files", []):
        path_name = item["path"]
        if path_name.startswith("core/templates/core/") and path_name.endswith(".html"):
            paths.add(path_name)
    views_content = read_project_or_generated(data, "core/views.py")
    for template_name in re.findall(r"template_name\s*=\s*[\"']([^\"']+)[\"']", views_content):
        if template_name.startswith("core/") and template_name.endswith(".html"):
            paths.add(f"core/templates/{template_name}")
    for path_name in ("core/templates/core/base.html", "core/templates/core/login.html"):
        if (PROJECT_ROOT / path_name).exists():
            paths.add(path_name)
    return paths


def parse_python(content, path_name, errors):
    if not content.strip():
        return None
    try:
        return ast.parse(content, filename=path_name)
    except SyntaxError as error:
        errors.append(f"{path_name}: синтаксическая ошибка Python: {error.msg}, строка {error.lineno}")
        return None


def call_name(node):
    if isinstance(node, ast.Call):
        function = node.func
        if isinstance(function, ast.Attribute):
            return function.attr
        if isinstance(function, ast.Name):
            return function.id
    return ""


def extract_python_names(tree):
    if tree is None:
        return set()
    names = set()
    for node in tree.body:
        if isinstance(node, (ast.ClassDef, ast.FunctionDef)):
            names.add(node.name)
    return names


def extract_model_fields(tree):
    fields = {}
    if tree is None:
        return fields
    for node in tree.body:
        if not isinstance(node, ast.ClassDef):
            continue
        current = {"id"}
        base_names = {base.id for base in node.bases if isinstance(base, ast.Name)}
        if "AbstractUser" in base_names:
            current.update(
                {
                    "password",
                    "last_login",
                    "is_superuser",
                    "username",
                    "first_name",
                    "last_name",
                    "email",
                    "is_staff",
                    "is_active",
                    "date_joined",
                    "groups",
                    "user_permissions",
                }
            )
        for statement in node.body:
            if isinstance(statement, ast.Assign) and call_name(statement.value) in MODEL_FIELD_TYPES:
                for target in statement.targets:
                    if isinstance(target, ast.Name):
                        current.add(target.id)
            if isinstance(statement, ast.AnnAssign) and isinstance(statement.target, ast.Name):
                current.add(statement.target.id)
        if len(current) > 1:
            fields[node.name] = current
    return fields


def extract_form_models_and_fields(tree):
    forms = []
    if tree is None:
        return forms
    for node in ast.walk(tree):
        if not isinstance(node, ast.ClassDef):
            continue
        declared_form_fields = set()
        for statement in node.body:
            if isinstance(statement, ast.Assign) and call_name(statement.value) in MODEL_FIELD_TYPES:
                for target in statement.targets:
                    if isinstance(target, ast.Name):
                        declared_form_fields.add(target.id)
        for child in node.body:
            if not isinstance(child, ast.ClassDef) or child.name != "Meta":
                continue
            model_name = ""
            fields = []
            for statement in child.body:
                if not isinstance(statement, ast.Assign):
                    continue
                targets = [target.id for target in statement.targets if isinstance(target, ast.Name)]
                if "model" in targets and isinstance(statement.value, ast.Name):
                    model_name = statement.value.id
                if "fields" in targets and isinstance(statement.value, (ast.List, ast.Tuple)):
                    fields = [item.value for item in statement.value.elts if isinstance(item, ast.Constant) and isinstance(item.value, str)]
                    fields = [field for field in fields if field not in declared_form_fields]
            if model_name and fields:
                forms.append((node.name, model_name, fields))
    return forms


def extract_update_or_create_fields(tree):
    writes = []
    if tree is None:
        return writes
    for node in ast.walk(tree):
        if not isinstance(node, ast.Call):
            continue
        function = node.func
        if not isinstance(function, ast.Attribute) or function.attr != "update_or_create":
            continue
        owner = function.value
        if not isinstance(owner, ast.Attribute) or owner.attr != "objects":
            continue
        if not isinstance(owner.value, ast.Name):
            continue
        model_name = owner.value.id
        defaults = next((keyword.value for keyword in node.keywords if keyword.arg == "defaults"), None)
        if isinstance(defaults, ast.Dict):
            fields = [key.value for key in defaults.keys if isinstance(key, ast.Constant) and isinstance(key.value, str)]
            writes.append((model_name, fields))
    return writes


def field_has_truthy_keyword(call, name):
    for keyword in call.keywords:
        if keyword.arg == name and isinstance(keyword.value, ast.Constant):
            return keyword.value.value is True
    return False


def field_has_keyword(call, name):
    return any(keyword.arg == name for keyword in call.keywords)


def extract_required_model_fields(tree):
    fields = {}
    if tree is None:
        return fields
    for node in tree.body:
        if not isinstance(node, ast.ClassDef):
            continue
        current = set()
        for statement in node.body:
            if not isinstance(statement, ast.Assign):
                continue
            call = statement.value
            field_type = call_name(call)
            if field_type not in MODEL_FIELD_TYPES:
                continue
            if field_type in ("AutoField", "BigAutoField"):
                continue
            if field_has_truthy_keyword(call, "null") or field_has_truthy_keyword(call, "blank") or field_has_keyword(call, "default"):
                continue
            for target in statement.targets:
                if isinstance(target, ast.Name):
                    current.add(target.id)
        fields[node.name] = current
    return fields


def extract_get_or_create_fields(tree):
    writes = []
    if tree is None:
        return writes
    for node in ast.walk(tree):
        if not isinstance(node, ast.Call):
            continue
        function = node.func
        if not isinstance(function, ast.Attribute) or function.attr != "get_or_create":
            continue
        owner = function.value
        if not isinstance(owner, ast.Attribute) or owner.attr != "objects":
            continue
        if not isinstance(owner.value, ast.Name):
            continue
        model_name = owner.value.id
        fields = {keyword.arg for keyword in node.keywords if keyword.arg and keyword.arg != "defaults"}
        defaults = next((keyword.value for keyword in node.keywords if keyword.arg == "defaults"), None)
        if isinstance(defaults, ast.Dict):
            fields.update(key.value for key in defaults.keys if isinstance(key, ast.Constant) and isinstance(key.value, str))
        writes.append((model_name, fields))
    return writes


def extract_string_constants(tree):
    if tree is None:
        return []
    return [node.value for node in ast.walk(tree) if isinstance(node, ast.Constant) and isinstance(node.value, str)]


def has_custom_user_model(models_content):
    return bool(re.search(r"class\s+User\s*\([^)]*AbstractUser[^)]*\)", models_content))


def auth_user_model_value(settings_content):
    match = re.search(r"\bAUTH_USER_MODEL\s*=\s*[\"']([^\"']+)[\"']", settings_content)
    return match.group(1) if match else ""


def has_list_page(views_content, urls_content):
    return "ListView" in views_content or bool(re.search(r"name\s*=\s*[\"'][^\"']*list[^\"']*[\"']", urls_content))


def views_have_list_controls(views_content):
    required_get = "request.GET" in views_content or "self.request.GET" in views_content
    has_filter = ".filter(" in views_content or " Q(" in views_content or "Q(" in views_content
    has_sort = ".order_by(" in views_content or "sort" in views_content or "order" in views_content
    return required_get and has_filter and has_sort


def is_list_template(path_name):
    normalized = normalize_relative_path(path_name)
    name = Path(normalized).stem
    if name in {"order_list"}:
        return False
    return normalized.startswith("core/templates/core/") and normalized.endswith(".html") and "list" in name


def template_has_list_controls(content):
    lower = content.lower()
    has_get_form = re.search(r"<form[^>]*method=[\"']get[\"']", content, flags=re.IGNORECASE) is not None
    has_search = 'name="q"' in lower or "name='q'" in lower or "search" in lower or "поиск" in lower
    has_filter = "filter" in lower or "фильтр" in lower or "name=\"type\"" in lower or "name='type'" in lower or "name=\"status\"" in lower or "name='status'" in lower
    has_sort = "sort" in lower or "сорт" in lower or "order" in lower
    return has_get_form and has_search and has_filter and has_sort


def extract_url_names_from_urls(content):
    return set(re.findall(r"name\s*=\s*[\"']([^\"']+)[\"']", content))


def extract_url_names_from_templates(content):
    return set(re.findall(r"{%\s*url\s+[\"']([^\"']+)[\"']", content))


def extract_url_names_from_python(content):
    patterns = [
        r"reverse_lazy\(\s*[\"']([^\"']+)[\"']",
        r"reverse\(\s*[\"']([^\"']+)[\"']",
        r"redirect\(\s*[\"']([^\"']+)[\"']",
    ]
    result = set()
    for pattern in patterns:
        result.update(re.findall(pattern, content))
    return result


def check_generated_consistency(data):
    errors = []
    python_trees = {}
    for path_name in CORE_GENERATION_FILES:
        content = read_project_or_generated(data, path_name)
        if content:
            python_trees[path_name] = parse_python(content, path_name, errors)

    urls_content = read_project_or_generated(data, "config/urls.py")
    views_content = read_project_or_generated(data, "core/views.py")
    models_content = read_project_or_generated(data, "core/models.py")
    models_tree = python_trees.get("core/models.py")
    forms_tree = python_trees.get("core/forms.py")
    views_tree = python_trees.get("core/views.py")
    import_tree = python_trees.get("core/management/commands/import_data.py")

    view_names = extract_python_names(views_tree)
    urls_tree = python_trees.get("config/urls.py")
    if urls_tree is not None:
        for node in ast.walk(urls_tree):
            if isinstance(node, ast.ImportFrom) and node.module == "core.views":
                for alias in node.names:
                    if alias.name not in view_names:
                        errors.append(f"config/urls.py импортирует {alias.name}, но такого класса/функции нет в core/views.py")

    url_names = extract_url_names_from_urls(urls_content)
    settings_content = read_project_or_generated(data, "config/settings.py")
    if "django.contrib.admin" in settings_content and "admin.site.urls" not in urls_content:
        errors.append("config/urls.py не содержит маршрут admin/. Нужно сохранить path('admin/', admin.site.urls).")
    if "name='partner_list'" in urls_content or 'name="partner_list"' in urls_content or re.search(r"name\s*=\s*[\"'][^\"']*list[^\"']*[\"']", urls_content):
        if not re.search(r"path\(\s*[\"']{2}\s*,", urls_content):
            errors.append("config/urls.py не содержит главную страницу ''. Добавь корневой маршрут на основной список.")
    for path_name in generated_template_paths(data):
        content = read_project_or_generated(data, path_name)
        if "```" in content:
            errors.append(f"{path_name} содержит markdown-ограждение ```, в HTML-файле его быть не должно")
        if "{% static" in content and "{% load static %}" not in content:
            errors.append(f"{path_name} использует static, но не содержит {{% load static %}}")
        if re.search(r"<form[^>]*method=[\"']post[\"']", content, flags=re.IGNORECASE) and "{% csrf_token %}" not in content:
            errors.append(f"{path_name} содержит POST-форму без {{% csrf_token %}}")
        for name in extract_url_names_from_templates(content):
            if name not in url_names:
                errors.append(f"{path_name} использует url '{name}', но такого name нет в config/urls.py")

    for name in extract_url_names_from_python(views_content):
        if name not in url_names:
            errors.append(f"core/views.py использует url '{name}', но такого name нет в config/urls.py")

    template_paths = generated_template_paths(data)
    for template_name in re.findall(r"template_name\s*=\s*[\"']core/([^\"']+)[\"']", views_content):
        path_name = f"core/templates/core/{template_name}"
        if path_name not in template_paths and not (PROJECT_ROOT / path_name).exists():
            errors.append(f"core/views.py ссылается на шаблон {path_name}, но файла нет")

    model_fields = extract_model_fields(models_tree)
    model_names = set(model_fields)
    form_names = extract_python_names(forms_tree)
    for path_name, tree in python_trees.items():
        if tree is None:
            continue
        for node in ast.walk(tree):
            if not isinstance(node, ast.ImportFrom):
                continue
            if node.module == "models" and node.level == 1:
                for alias in node.names:
                    if alias.name not in model_names:
                        errors.append(f"{path_name} импортирует модель {alias.name}, но ее нет в core/models.py")
            if node.module == "forms" and node.level == 1:
                for alias in node.names:
                    if alias.name not in form_names:
                        errors.append(f"{path_name} импортирует форму {alias.name}, но ее нет в core/forms.py")

    for form_name, model_name, fields in extract_form_models_and_fields(forms_tree):
        if model_name not in model_fields:
            errors.append(f"{form_name} использует модель {model_name}, но она не найдена в core/models.py")
            continue
        for field in fields:
            if field not in model_fields[model_name]:
                errors.append(f"{form_name} использует поле {model_name}.{field}, но такого поля нет в core/models.py")

    for model_name, fields in extract_update_or_create_fields(import_tree):
        if model_name not in model_fields:
            continue
        for field in fields:
            if field not in model_fields[model_name]:
                errors.append(f"import_data.py записывает поле {model_name}.{field}, но такого поля нет в core/models.py")

    required_fields = extract_required_model_fields(models_tree)
    for model_name, fields in extract_get_or_create_fields(import_tree):
        missing = sorted(required_fields.get(model_name, set()) - set(fields))
        if missing:
            errors.append(f"import_data.py вызывает {model_name}.objects.get_or_create без обязательных полей: {', '.join(missing)}. Для уже импортированных связей используй get(), а не get_or_create().")

    for value in extract_string_constants(import_tree):
        lower = value.lower()
        if not any(lower.endswith(suffix) for suffix in IMPORT_EXTENSIONS):
            continue
        normalized = value.replace("\\", "/")
        if normalized.startswith("Ресурсы/") or normalized.startswith("tools/input/") or "/" in normalized:
            errors.append(f"import_data.py использует путь '{value}'. Excel/CSV должны читаться из core/import через BASE_DIR, без папки Ресурсы/tools/input.")

    if has_list_page(views_content, urls_content) and not views_have_list_controls(views_content):
        errors.append("core/views.py содержит list-страницу без полноценного get_queryset() с поиском, фильтрацией и сортировкой через GET-параметры")

    for path_name in generated_template_paths(data):
        if is_list_template(path_name):
            content = read_project_or_generated(data, path_name)
            if not template_has_list_controls(content):
                errors.append(f"{path_name} является list-шаблоном, но не содержит GET-форму поиска, фильтра и сортировки")

    for path_name in generated_template_paths(data):
        content = read_project_or_generated(data, path_name)
        if re.search(r"{%\s*extends\s+[\"']base\.html[\"']\s*%}", content):
            errors.append(f"{path_name} расширяет base.html, нужно использовать core/base.html")

    return errors


def check_schema_consistency(data):
    errors = []
    python_trees = {}
    for path_name in SCHEMA_STAGE_FILES:
        content = read_project_or_generated(data, path_name)
        if content:
            python_trees[path_name] = parse_python(content, path_name, errors)

    models_tree = python_trees.get("core/models.py")
    import_tree = python_trees.get("core/management/commands/import_data.py")
    models_content = read_project_or_generated(data, "core/models.py")
    settings_content = read_project_or_generated(data, "config/settings.py")
    model_fields = extract_model_fields(models_tree)
    model_names = set(model_fields)

    if auth_user_model_value(settings_content) == "core.User" and not has_custom_user_model(models_content):
        errors.append("config/settings.py требует AUTH_USER_MODEL = core.User, но core/models.py не содержит class User(AbstractUser). Нельзя удалять User из моделей.")

    for path_name, tree in python_trees.items():
        if tree is None:
            continue
        for node in ast.walk(tree):
            if not isinstance(node, ast.ImportFrom):
                continue
            if node.module == "models" and node.level == 1:
                for alias in node.names:
                    if alias.name not in model_names:
                        errors.append(f"{path_name} импортирует модель {alias.name}, но ее нет в core/models.py")

    for model_name, fields in extract_update_or_create_fields(import_tree):
        if model_name not in model_fields:
            continue
        for field in fields:
            if field not in model_fields[model_name]:
                errors.append(f"import_data.py записывает поле {model_name}.{field}, но такого поля нет в core/models.py")

    required_fields = extract_required_model_fields(models_tree)
    for model_name, fields in extract_get_or_create_fields(import_tree):
        missing = sorted(required_fields.get(model_name, set()) - set(fields))
        if missing:
            errors.append(f"import_data.py вызывает {model_name}.objects.get_or_create без обязательных полей: {', '.join(missing)}. Для уже импортированных связей используй get(), а не get_or_create().")

    for value in extract_string_constants(import_tree):
        lower = value.lower()
        if not any(lower.endswith(suffix) for suffix in IMPORT_EXTENSIONS):
            continue
        normalized = value.replace("\\", "/")
        if normalized.startswith("Ресурсы/") or normalized.startswith("tools/input/") or "/" in normalized:
            errors.append(f"import_data.py использует путь '{value}'. Excel/CSV должны читаться из core/import через BASE_DIR, без папки Ресурсы/tools/input.")

    return errors


def manifest_from_paths(summary, paths, reason):
    return {
        "summary": summary,
        "steps": [],
        "commands": [],
        "files": [{"path": path_name, "reason": reason} for path_name in paths],
        "assets": [],
    }


def normalize_template_file_path(template):
    normalized = normalize_relative_path(template)
    if not normalized or not normalized.endswith(".html"):
        return ""
    if normalized.startswith("core/templates/core/"):
        return normalized
    if normalized.startswith("core/"):
        return f"core/templates/{normalized}"
    if "/" not in normalized:
        return f"core/templates/core/{normalized}"
    return normalized


def blueprint_templates(blueprint_text):
    try:
        blueprint = json.loads(blueprint_text)
    except json.JSONDecodeError:
        return []
    templates = []
    for page in blueprint.get("pages", []):
        template = normalize_template_file_path(page.get("template", ""))
        if template.startswith("core/templates/core/") and template.endswith(".html"):
            templates.append(template)
    return templates


def blueprint_assets(blueprint_text):
    try:
        blueprint = json.loads(blueprint_text)
    except json.JSONDecodeError:
        return []
    return blueprint.get("assets", [])


def get_or_create_blueprint():
    path = OUTPUT_DIR / "blueprint.md"
    if path.exists():
        return read_text(path)
    assignment_context = collect_assignment()
    return get_blueprint(assignment_context)


def build_stage_change_set(context, manifest, check_func, stage_name):
    data = build_full_change_set(context, manifest)
    last_errors = []
    for attempt in range(1, MAX_REPAIR_ATTEMPTS + 1):
        errors = check_func(data)
        if not errors:
            return data, []
        last_errors = errors
        errors_text = "\n".join(f"- {error}" for error in errors)
        save_output(f"{stage_name}_errors_{attempt}.md", errors_text)
        log(f"{stage_name}: проверка нашла проблем: {len(errors)}. Перегенерирую файлы {attempt}/{MAX_REPAIR_ATTEMPTS}...", 82)
        runtime_context = build_runtime_context(context, data["files"], errors_text)
        data = build_full_change_set(runtime_context, manifest)

    final_errors = check_func(data)
    return data, final_errors or last_errors


def validate_schema_project():
    command = ["uv", "run", "python", "-m", "py_compile", *SCHEMA_STAGE_FILES]
    command_text = " ".join(command)
    log(f"Проверяю схему: {command_text}")
    ok, output = run_project_command(command)
    return ok, f"$ {command_text}\n{output or 'OK'}"


def replace_generated_files(data, fixes):
    current = {item["path"]: item for item in data.get("files", [])}
    for item in fixes.get("files", []):
        current[item["path"]] = item
    data["files"] = sort_file_items(list(current.values()))
    return data


def expand_manifest_file_items(context, files):
    expanded = []
    for item in filter_generated_files(files):
        path_name = normalize_relative_path(item.get("path", ""))
        reason = item.get("reason", "")
        if not path_name:
            continue

        if looks_like_folder(path_name):
            log(f"Уточняю файлы внутри папки: {path_name}")
            answer = call_yandex(build_folder_files_prompt(context, path_name, reason))
            save_output("last_response.md", answer)
            save_output(f"response_folder_{safe_output_name(path_name)}.md", answer)
            data = extract_json(answer)
            for child in filter_generated_files(data.get("files", [])):
                child_path = validate_file_path(child.get("path", ""))
                child_reason = child.get("reason") or reason
                expanded.append({"path": child_path, "reason": child_reason})
            continue

        path_name = validate_file_path(path_name)
        expanded.append({"path": path_name, "reason": reason})

    unique = []
    seen = set()
    for item in expanded:
        if item["path"] in seen:
            continue
        seen.add(item["path"])
        unique.append(item)
    return unique


def build_full_change_set(context, manifest):
    result = dict(manifest)
    result["files"] = []
    files = sort_file_items(filter_generated_files(manifest.get("files", [])))
    total = len(files)
    order_note = "Файлы текущего apply в порядке генерации:\n" + "\n".join(f"- {item['path']}" for item in files)
    for index, item in enumerate(files, start=1):
        path_name = item.get("path", "")
        reason = item.get("reason", "")
        if not path_name:
            continue
        path_name = validate_file_path(path_name)
        percent = 30 + round(index / max(total, 1) * 50)
        log(f"генерирует файл {index}/{total}: {path_name}", percent)
        runtime_context = build_file_context(context, result["files"], path_name, order_note)
        answer = ""
        content = None
        last_error = None
        for attempt in range(1, MAX_FILE_GENERATION_ATTEMPTS + 1):
            if attempt == 1:
                prompt = build_file_prompt(runtime_context, path_name, reason)
                output_name = f"response_{safe_output_name(path_name)}.md"
            else:
                log(f"Повторяю генерацию файла {attempt}/{MAX_FILE_GENERATION_ATTEMPTS}: {path_name}")
                retry_context = build_file_context(
                    context,
                    result["files"],
                    path_name,
                    order_note + f"\nПредыдущая ошибка генерации файла: {last_error}",
                )
                prompt = build_file_retry_prompt(retry_context, path_name, reason, answer)
                output_name = f"response_retry_{attempt}_{safe_output_name(path_name)}.md"

            answer = call_yandex(prompt)
            save_output("last_response.md", answer)
            save_output(output_name, answer)
            try:
                content = extract_file_content(answer, path_name)
                break
            except ValueError as error:
                last_error = str(error)
                if attempt == MAX_FILE_GENERATION_ATTEMPTS:
                    raise

        if content is None:
            raise ValueError(f"Не удалось сгенерировать файл {path_name}")
        content = clean_generated_content(path_name, content)
        result["files"].append({"path": path_name, "content": content})
    return result


def repair_generated_consistency(context, data, attempts=MAX_REPAIR_ATTEMPTS):
    last_errors = []
    for attempt in range(1, attempts + 1):
        errors = check_generated_consistency(data)
        if not errors:
            return data, []

        last_errors = errors
        errors_text = "\n".join(f"- {error}" for error in errors)
        save_output(f"consistency_errors_{attempt}.md", errors_text)
        log(f"Внутренняя проверка нашла проблем: {len(errors)}. Исправляю до записи файлов...", 82)

        runtime_context = build_runtime_context(context, data["files"], errors_text)
        answer = call_yandex(build_consistency_repair_prompt(runtime_context, errors_text))
        save_output("last_response.md", answer)
        save_output(f"consistency_repair_manifest_{attempt}.md", answer)
        manifest = extract_json(answer)
        manifest["files"] = filter_generated_files(manifest.get("files", []))
        manifest["files"] = expand_manifest_file_items(runtime_context, manifest.get("files", []))
        if not manifest["files"]:
            break
        fixes = build_full_change_set(runtime_context, manifest)
        data = replace_generated_files(data, fixes)

    final_errors = check_generated_consistency(data)
    return data, final_errors or last_errors


def repair_project(error_text, attempts=MAX_REPAIR_ATTEMPTS):
    all_changed = []
    last_error = error_text
    last_backup = None

    for attempt in range(1, attempts + 1):
        log(f"Исправление после проверки, попытка {attempt}/{attempts}...", 94)
        context = build_context()
        answer = call_yandex(build_repair_prompt(context, last_error))
        save_output("last_response.md", answer)
        save_output(f"repair_manifest_{attempt}.md", answer)

        manifest = extract_json(answer)
        manifest["files"] = filter_generated_files(manifest.get("files", []))
        manifest["files"] = expand_manifest_file_items(context, manifest.get("files", []))
        if not manifest["files"]:
            raise ValueError("Проверка упала, но список файлов для исправления пустой.")

        data = build_full_change_set(context, manifest)
        changed, backup = write_files(data)
        all_changed.extend(changed)
        last_backup = backup

        ok, check_output = validate_project()
        save_output(f"repair_check_{attempt}.md", check_output)
        if ok:
            return True, check_output, all_changed, last_backup

        last_error = check_output

    return False, last_error, all_changed, last_backup


def save_output(name, content):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / name
    path.write_text(content, encoding="utf-8")
    return path


def save_answer(question, answer):
    timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    text = f"# Вопрос\n\n{question}\n\n# Ответ\n\n{answer}\n"
    path = save_output(f"answer_{timestamp}.md", text)
    chat_log = OUTPUT_DIR / "chat_log.md"
    with chat_log.open("a", encoding="utf-8") as file:
        file.write(f"\n\n## {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        file.write(text)
    return path


def command_collect():
    log("Собираю контекст из input и файлов проекта...", 10)
    context = build_context()
    path = save_output("context_preview.md", context)
    assignment_count = len(collect_input_files())
    log(f"Контекст собран. Файлов задания: {assignment_count}.", 100)
    print(f"Файл: {path}")


def command_plan():
    log("Собираю контекст...", 10)
    context = build_context()
    prompt = build_plan_prompt(context)
    log(" анализирует задание и готовит план...", 40)
    answer = call_yandex(prompt)
    save_output("last_response.md", answer)
    path = save_output("plan.md", answer)
    log("План сохранен.", 100)
    print(f"Файл: {path}")


def command_schema():
    log("Собираю задание и Excel...", 5)
    assignment_context = collect_assignment()
    log("строит blueprint для моделей и импорта...", 10)
    blueprint = get_blueprint(assignment_context)
    context = build_generation_context(blueprint)
    context = "\n".join(
        [
            context,
            "\n# РЕЖИМ SCHEMA",
            "Сейчас нужно сгенерировать только модели, admin.py и import_data.py.",
            "Не трогай forms.py, views.py, urls.py, templates, css и settings.py.",
            "core/models.py должен быть похож по простоте на текущий учебный Django-код, но полностью под новые Excel-данные.",
            "Если settings.py содержит AUTH_USER_MODEL = 'core.User', core/models.py обязан содержать class User(AbstractUser).",
            "import_data.py должен быть простым, идемпотентным и читать Excel/CSV только из core/import.",
            "core/admin.py должен регистрировать только реальные модели из core/models.py.",
        ]
    )
    manifest = manifest_from_paths(
        "Схема данных и импорт",
        SCHEMA_STAGE_FILES,
        "этап schema: модели, регистрация в админке и импорт данных",
    )
    data, errors = build_stage_change_set(context, manifest, check_schema_consistency, "schema")
    if errors:
        errors_text = "\n".join(f"- {error}" for error in errors)
        save_output("schema_errors_final.md", errors_text)
        raise ValueError("schema не прошел внутреннюю проверку. Подробности: tools/output/schema_errors_final.md")

    log("Записываю schema-файлы и создаю backup...", 85)
    changed, backup = write_files(data)
    log("Копирую Excel/CSV для import_data...", 92)
    copied_imports = copy_import_resources(backup)
    check_ok, check_output = validate_schema_project()
    save_output("schema_check.md", check_output)

    report = [
        f"# Отчет schema {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Измененные файлы",
        *(f"- {path}" for path in changed),
        "",
        "## Скопированные файлы импорта",
        *(f"- {path}" for path in copied_imports),
        "",
        "## Backup",
        str(backup),
        "",
        "## Проверка",
        "OK" if check_ok else "ОШИБКА",
        "",
        "```text",
        check_output,
        "```",
    ]
    path = save_output("schema_report.md", "\n".join(report))
    log("Schema готов.", 100)
    print(f"Изменено файлов: {len(changed)}")
    print(f"Скопировано файлов импорта: {len(copied_imports)}")
    print(f"Проверка schema: {'OK' if check_ok else 'ОШИБКА'}")
    print(f"Отчет: {path}")
    print(f"Backup: {backup}")
    if not check_ok:
        sys.exit(1)


def command_interface():
    log("Собираю blueprint и текущие schema-файлы...", 5)
    blueprint = get_or_create_blueprint()
    context = build_generation_context(blueprint)
    context = "\n".join(
        [
            context,
            "\n# РЕЖИМ INTERFACE",
            "Сейчас нужно сгенерировать только интерфейс: forms.py, views.py, urls.py, templates, css, permissions/context_processors.",
            "Не трогай core/models.py, core/admin.py, import_data.py и config/settings.py.",
            "Бери модели и поля из текущего core/models.py как источник правды.",
            "Сохрани все страницы из BLUEPRINT: список, добавление, редактирование, просмотр, история, если они там есть.",
            "Главная list-страница должна иметь поиск, фильтр и сортировку через GET.",
            "config/urls.py должен содержать все нужные маршруты, включая главную страницу и login/logout при наличии login.html/base.html.",
        ]
    )
    paths = list(INTERFACE_STAGE_BASE_FILES)
    for template in blueprint_templates(blueprint):
        if template not in paths:
            paths.append(template)
    manifest = manifest_from_paths(
        "Интерфейс проекта",
        paths,
        "этап interface: формы, views, urls, templates, css и права интерфейса",
    )
    data, errors = build_stage_change_set(context, manifest, check_generated_consistency, "interface")
    if errors:
        errors_text = "\n".join(f"- {error}" for error in errors)
        save_output("interface_errors_final.md", errors_text)
        raise ValueError("interface не прошел внутреннюю проверку. Подробности: tools/output/interface_errors_final.md")

    assets = validate_assets({"assets": blueprint_assets(blueprint)})
    log("Записываю interface-файлы и создаю backup...", 85)
    changed, backup = write_files(data)
    log("Копирую картинки...", 92)
    copied_assets = copy_assets(assets, backup)
    log("Проверяю проект после interface...", 93)
    check_ok, check_output = validate_project()
    save_output("interface_check.md", check_output)
    repair_changed = []
    repair_backup = None
    if not check_ok:
        log("Interface-проверка упала, запускаю repair...", 94)
        check_ok, check_output, repair_changed, repair_backup = repair_project(check_output)
        save_output("interface_check_after_repair.md", check_output)
        changed.extend(repair_changed)

    report = [
        f"# Отчет interface {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Измененные файлы",
        *(f"- {path}" for path in changed),
        "",
        "## Скопированные картинки",
        *(f"- {path}" for path in copied_assets),
        "",
        "## Backup",
        str(backup),
        "",
        "## Backup исправления",
        str(repair_backup or ""),
        "",
        "## Проверка",
        "OK" if check_ok else "ОШИБКА",
        "",
        "```text",
        check_output,
        "```",
    ]
    path = save_output("interface_report.md", "\n".join(report))
    log("Interface готов.", 100)
    print(f"Изменено файлов: {len(changed)}")
    print(f"Скопировано картинок: {len(copied_assets)}")
    print(f"Проверка interface: {'OK' if check_ok else 'ОШИБКА'}")
    print(f"Отчет: {path}")
    print(f"Backup: {backup}")
    if repair_backup:
        print(f"Backup исправления: {repair_backup}")
    if not check_ok:
        sys.exit(1)


def command_audit(condition):
    log("Собираю контекст для аудита...", 10)
    context = build_context()
    log(" проверяет соответствие требованиям...", 50)
    answer = call_yandex(build_audit_prompt(context, condition))
    save_output("last_response.md", answer)
    path = save_output("audit.md", answer)
    log("Аудит сохранен.", 100)
    print(answer.strip())
    print(f"\nОтчет: {path}")


def command_apply():
    log("Собираю контекст...", 5)
    assignment_context = collect_assignment()
    log("строит единый blueprint проекта...", 10)
    blueprint = get_blueprint(assignment_context)
    context = build_generation_context(blueprint)
    prompt = build_apply_prompt(context)
    log(" выбирает список файлов для изменения...", 15)
    answer = call_yandex(prompt)
    save_output("last_response.md", answer)
    save_output("apply_manifest_response.md", answer)
    log("Разбираю список изменений от AI...", 25)
    manifest = extract_json(answer)
    manifest = ensure_core_files(manifest)
    manifest = ensure_blueprint_templates(manifest, blueprint)
    manifest = ensure_support_files(manifest)
    assets = validate_assets(manifest)
    manifest["files"] = expand_manifest_file_items(context, manifest.get("files", []))
    files_count = len(manifest.get("files", []))
    log(f" предложил файлов: {files_count}, картинок: {len(assets)}.", 30)
    for source, _, target_name in assets:
        source_label = source.relative_to(INPUT_DIR).as_posix()
        log(f"Картинка найдена: {source_label} -> {target_name}")
    data = build_full_change_set(context, manifest)
    log("Проверяю согласованность сгенерированных файлов...", 82)
    data, consistency_errors = repair_generated_consistency(context, data)
    if consistency_errors:
        errors_text = "\n".join(f"- {error}" for error in consistency_errors)
        save_output("consistency_errors_final.md", errors_text)
        raise ValueError("Сгенерированные файлы все еще не согласованы. Подробности: tools/output/consistency_errors_final.md")
    log("Записываю файлы и создаю backup...", 85)
    changed, backup = write_files(data)
    log("Копирую картинки...", 92)
    copied_assets = copy_assets(assets, backup)
    log("Копирую Excel/CSV для import_data...", 92)
    copied_imports = copy_import_resources(backup)
    log("Проверяю проект после изменений...", 93)
    check_ok, check_output = validate_project()
    save_output("check_after_apply.md", check_output)
    repair_changed = []
    repair_backup = None
    if not check_ok:
        log("Проверка упала, запускаю автоматическое исправление...", 94)
        check_ok, check_output, repair_changed, repair_backup = repair_project(check_output)
        save_output("check_after_repair.md", check_output)
        changed.extend(repair_changed)
    report = [
        f"# Отчет AI-помощника {dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        data.get("summary", ""),
        "",
        "## Измененные файлы",
        *(f"- {path}" for path in changed),
        "",
        "## Скопированные картинки",
        *(f"- {path}" for path in copied_assets),
        "",
        "## Скопированные файлы импорта",
        *(f"- {path}" for path in copied_imports),
        "",
        "## Backup",
        str(backup),
        "",
        "## Backup исправления",
        str(repair_backup or ""),
        "",
        "## Проверка Django",
        "OK" if check_ok else "ОШИБКА",
        "",
        "```text",
        check_output,
        "```",
        "",
        "## Команды",
        *(f"```powershell\n{command}\n```" for command in data.get("commands", [])),
        "",
        "## Проверить",
        *(f"- {step}" for step in data.get("steps", [])),
    ]
    path = save_output("apply_report.md", "\n".join(report))
    log("Готово.", 100)
    print(f"Изменено файлов: {len(changed)}")
    print(f"Скопировано картинок: {len(copied_assets)}")
    print(f"Скопировано файлов импорта: {len(copied_imports)}")
    print(f"Проверка Django: {'OK' if check_ok else 'ОШИБКА'}")
    print(f"Отчет: {path}")
    print(f"Backup: {backup}")
    if repair_backup:
        print(f"Backup исправления: {repair_backup}")
    if not check_ok:
        sys.exit(1)


def command_test():
    log("Проверяю подключение ...", 50)
    answer = call_yandex("Ответь одним предложением: API Yandex AI Studio работает.")
    log("Ответ получен.", 100)
    print(answer.strip())


def command_ask(question):
    if not question:
        raise ValueError('Напиши вопрос после команды, например: uv run python tool.py ask "почему migrate падает?"')
    log("Собираю контекст для вопроса...", 10)
    context = build_context()
    log("Отправляю вопрос в AI...", 50)
    answer = call_yandex(build_question_prompt(context, question, []))
    path = save_answer(question, answer)
    log("Ответ получен.", 100)
    print()
    print(answer.strip())
    print()
    print(f"Ответ сохранен: {path}")


def command_chat():
    log("Подготавливаю чат...", 10)
    history = []
    log("Чат готов. Напиши вопрос или exit для выхода.", 100)
    while True:
        question = input("\nТы> ").strip()
        if question.lower() in ("exit", "quit", "выход"):
            print("Чат завершен.")
            return
        if not question:
            continue
        log("Собираю свежий контекст проекта...", 20)
        context = build_context()
        log(" думает...", 50)
        answer = call_yandex(build_question_prompt(context, question, history))
        history.append(f"Пользователь: {question}")
        history.append(f": {answer}")
        save_answer(question, answer)
        print(f"\nAI> {answer.strip()}")


def find_backup(name=""):
    if name:
        backup = (BACKUP_DIR / name).resolve()
        if not is_inside(BACKUP_DIR, backup) or not backup.exists() or not backup.is_dir():
            raise ValueError(f"Backup не найден: {name}")
        return backup

    backups = sorted((path for path in BACKUP_DIR.iterdir() if path.is_dir()), reverse=True)
    if not backups:
        raise ValueError("Backup пока нет.")
    return backups[0]


def command_restore(name=""):
    backup = find_backup(name)
    restored = []
    log(f"Восстанавливаю из backup: {backup.name}", 20)
    for source in backup.rglob("*"):
        if not source.is_file():
            continue
        relative = source.relative_to(backup).as_posix()
        target = (PROJECT_ROOT / relative).resolve()
        if not is_inside(PROJECT_ROOT, target):
            raise ValueError(f"Путь вышел за пределы проекта: {relative}")
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)
        restored.append(relative)
    log("Восстановление завершено.", 100)
    print(f"Backup: {backup}")
    print(f"Восстановлено файлов: {len(restored)}")
    for path in restored:
        print(f"- {path}")


def command_check():
    ok, output = validate_project()
    path = save_output("manual_check.md", output)
    print(output)
    print(f"\nОтчет: {path}")
    if not ok:
        sys.exit(1)


def command_test_pages():
    ok, output = validate_pages()
    save_output("pages_check_before_repair.md", output)
    changed = []
    backup = None
    attempt = 0

    while not ok and not page_test_waits_for_database(output) and attempt < MAX_REPAIR_ATTEMPTS:
        attempt += 1
        log(f"Проверка страниц упала, исправление {attempt}/{MAX_REPAIR_ATTEMPTS}...", 94)
        _, repair_output, repair_changed, backup = repair_project(output, attempts=1)
        changed.extend(repair_changed)
        ok, output = validate_pages()
        save_output(f"pages_check_after_repair_{attempt}.md", output)
        if not repair_changed and repair_output == output:
            break

    path = save_output("pages_check.md", output)
    print(output)
    print(f"\nОтчет: {path}")
    if changed:
        print(f"Изменено файлов: {len(changed)}")
    if backup:
        print(f"Backup исправления: {backup}")
    if not ok:
        sys.exit(1)


def command_repair():
    ok, output = validate_project()
    save_output("manual_check_before_repair.md", output)
    if ok:
        print("Проверка Django: OK")
        return
    ok, output, changed, backup = repair_project(output)
    save_output("manual_check_after_repair.md", output)
    print(f"Изменено файлов: {len(changed)}")
    print(f"Backup исправления: {backup}")
    print(f"Проверка Django: {'OK' if ok else 'ОШИБКА'}")
    if not ok:
        print(output)
        sys.exit(1)


def main():
    ensure_dirs()
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=["collect", "test", "plan", "schema", "interface", "apply", "ask", "chat", "audit", "restore", "check", "test_pages", "repair"])
    parser.add_argument("question", nargs=argparse.REMAINDER)
    args = parser.parse_args()

    commands = {
        "collect": command_collect,
        "test": command_test,
        "plan": command_plan,
        "schema": command_schema,
        "interface": command_interface,
        "apply": command_apply,
        "ask": lambda: command_ask(" ".join(args.question).strip()),
        "chat": command_chat,
        "audit": lambda: command_audit(" ".join(args.question).strip()),
        "restore": lambda: command_restore(" ".join(args.question).strip()),
        "check": command_check,
        "test_pages": command_test_pages,
        "repair": command_repair,
    }
    try:
        commands[args.command]()
    except Exception as error:
        print(f"Ошибка: {error}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
